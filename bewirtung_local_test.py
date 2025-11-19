from docx import Document
from docx2pdf import convert
from pypdf import PdfReader, PdfWriter
from datetime import date
import os

# --- 1. Pfade definieren ---
TEMPLATE_PATH = "templates/bewirtung_template.docx"  # deine Vorlage
OUTPUT_DOCX = "output/bewirtung_fertig.docx"         # ausgefülltes Word-Dokument
OUTPUT_PDF_FORM = "output/bewirtung_fertig.pdf"      # PDF nur mit Formular
INPUT_RECEIPT_PDF = "input/bon_beispiel.pdf"         # dein Test-Bon
OUTPUT_FINAL_PDF = "output/bewirtung_final.pdf"      # Bon + Formular zusammen

# --- 2. Beispiel-Daten (später kommen die vom Bot/LLM) ---
# bewirtungsdatum = Tag des Essens (vom Bon)
# unterschriftsdatum = heute oder wann du unterschreibst
data = {
    "bewirtungsdatum": "09.07.2025",
    "unterschriftsdatum": date.today().strftime("%d.%m.%Y"),  # heute, z.B. 19.11.2025
    "ort": "Berlin",
    "restaurant": "SaPHI Sushi & Bowl",
    "adresse": "Reichenberger Str. 120, 10999 Berlin",
    "anlass": "Nachbesprechung ZuBerlin, Lessons Learned und Aufgabenverteilung.",
    "personen": [
        "Christian Haug, ZuBerlin",
        "Pascal Stichler, ZuBerlin",
        "Johannes Köhler, ZuBerlin",
    ],
    "betrag": "30,60 EUR"
}


def fill_template(template_path: str, output_path: str, data: dict) -> None:
    """
    Füllt die Word-Vorlage mit den Platzhaltern aus dem data-Dict.
    Erwartete Platzhalter im DOCX:
      {{bewirtungsdatum}}
      {{unterschriftsdatum}}
      {{ort}}
      {{restaurant}}
      {{adresse}}
      {{anlass}}
      {{personen}}
      {{betrag}}
    """
    doc = Document(template_path)

    placeholders = {
        "{{bewirtungsdatum}}": data["bewirtungsdatum"],
        "{{unterschriftsdatum}}": data["unterschriftsdatum"],
        "{{ort}}": data["ort"],
        "{{restaurant}}": data["restaurant"],
        "{{adresse}}": data["adresse"],
        "{{anlass}}": data["anlass"],
        "{{personen}}": "\n".join(data["personen"]),
        "{{betrag}}": data["betrag"],
    }

    # Durch alle Absätze gehen und Text ersetzen
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Tabellen (falls deine Vorlage Tabellen benutzt) ebenfalls durchsuchen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    doc.save(output_path)


def merge_pdfs(receipt_pdf_path: str, form_pdf_path: str, output_path: str) -> None:
    """
    Fügt zuerst alle Seiten des Bons und danach alle Seiten
    des Formular-PDFs zu einer gemeinsamen PDF zusammen.
    """
    writer = PdfWriter()

    # 1) Quittung
    receipt_reader = PdfReader(receipt_pdf_path)
    for page in receipt_reader.pages:
        writer.add_page(page)

    # 2) Formular
    form_reader = PdfReader(form_pdf_path)
    for page in form_reader.pages:
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)


def main():
    # 1) Prüfen, ob Vorlage existiert
    if not os.path.exists(TEMPLATE_PATH):
        print(f"Fehler: Vorlage nicht gefunden unter {TEMPLATE_PATH}")
        return

    # 2) Template ausfüllen -> DOCX
    fill_template(TEMPLATE_PATH, OUTPUT_DOCX, data)
    print("DOCX-Formular erzeugt:", OUTPUT_DOCX)

    # 3) DOCX -> PDF
    try:
        convert(OUTPUT_DOCX, OUTPUT_PDF_FORM)
        print("PDF-Formular erzeugt:", OUTPUT_PDF_FORM)
    except Exception as e:
        print("Fehler bei DOCX -> PDF Konvertierung:", e)
        print("Falls das auf dem Mac Probleme macht, können wir eine Alternative nutzen.")
        return

    # 4) Bon + Formular zusammenführen
    if not os.path.exists(INPUT_RECEIPT_PDF):
        print(f"Fehler: Bitte lege eine Bon-PDF in '{INPUT_RECEIPT_PDF}' ab.")
        return

    merge_pdfs(INPUT_RECEIPT_PDF, OUTPUT_PDF_FORM, OUTPUT_FINAL_PDF)
    print("Fertige Bewirtungs-PDF erzeugt:", OUTPUT_FINAL_PDF)


if __name__ == "__main__":
    main()
    