from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse
import os
import json

from docx import Document
from PyPDF2 import PdfReader, PdfWriter

from ocr_bon import ocr_bon
from extract_agent_gemini import extract_bewirtungsdaten_gemini
import subprocess
from pathlib import Path

def docx_to_pdf_libreoffice(input_docx: str, output_pdf: str) -> None:
    outdir = str(Path(output_pdf).parent)
    Path(outdir).mkdir(parents=True, exist_ok=True)

    # LibreOffice erzeugt PDF mit gleichem Dateinamen wie DOCX
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            input_docx,
            "--outdir",
            outdir,
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    produced = Path(outdir) / (Path(input_docx).stem + ".pdf")
    if produced != Path(output_pdf):
        produced.replace(output_pdf)
# --------------------------------------------------
# FastAPI App
# --------------------------------------------------

app = FastAPI()

# --------------------------------------------------
# Pfade / Konstanten
# --------------------------------------------------

TEMPLATE_PATH = "templates/bewirtung_template.docx"
OUTPUT_DOCX = "output/bewirtung_fertig.docx"
OUTPUT_PDF_FORM = "output/bewirtung_fertig.pdf"
OUTPUT_FINAL_PDF = "output/bewirtungs_beleg_final.pdf"


# --------------------------------------------------
# Template ausfüllen (Word -> DOCX)
# --------------------------------------------------

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Inches
import os

SIGNATURE_DIR = Path("signatures")
DEFAULT_SIGNATURES = [
    SIGNATURE_DIR / "default.png",
    SIGNATURE_DIR / "default.jpg",
    SIGNATURE_DIR / "default.jpeg",
]

def _get_default_signature_path() -> str | None:
    for p in DEFAULT_SIGNATURES:
        if p.exists():
            return str(p)
    return None


def fill_template(bew_data: dict):
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)

    # Unterschriftsdatum immer = heutiges Datum (Erstellungsdatum des Belegs)
    bew_data = bew_data.copy()
    bew_data["unterschriftsdatum"] = date.today().strftime("%d.%m.%Y")


    # Personenliste → String
    if "personen" in bew_data and isinstance(bew_data["personen"], list):
        bew_data = bew_data.copy()
        bew_data["personen"] = ", ".join(bew_data["personen"])

    doc = Document(TEMPLATE_PATH)
    signature_path = _get_default_signature_path()

    def replace_text(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text
        for key, val in bew_data.items():
            if key == "signature":
                continue
            new_text = new_text.replace(f"{{{{{key}}}}}", str(val))

        if new_text != full_text:
            paragraph.clear()
            paragraph.add_run(new_text)

    def replace_signature(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        if "{{signature}}" not in full_text:
            return

        paragraph.clear()

        if signature_path:
            run = paragraph.add_run()
            run.add_picture(signature_path, width=Inches(1.6))
        else:
            paragraph.add_run("(bitte unterschreiben)")

    # Normale Absätze
    for p in doc.paragraphs:
        if "{{" in p.text:
            replace_text(p)
        if "{{signature}}" in "".join(r.text for r in p.runs):
            replace_signature(p)

    # Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text:
                        replace_text(p)
                    if "{{signature}}" in "".join(r.text for r in p.runs):
                        replace_signature(p)

    doc.save(OUTPUT_DOCX)


# --------------------------------------------------
# DOCX -> PDF konvertieren
# --------------------------------------------------


def generate_form_pdf(bew_data: dict) -> None:
    """
    Nimmt Bewirtungsdaten, füllt das DOCX-Template
    und konvertiert danach in eine PDF (OUTPUT_PDF_FORM).
    """
    fill_template(bew_data)
    docx_to_pdf_libreoffice(OUTPUT_DOCX, OUTPUT_PDF_FORM)


import re
from decimal import Decimal, InvalidOperation

def _parse_eur_amount(s: str) -> Decimal | None:
    """
    Parses amounts like '34,00 EUR', '34.00', 'EUR 34,00' into Decimal(34.00).
    Returns None if not parseable.
    """
    if not s:
        return None
    s = str(s).strip()

    m = re.search(r'(\d{1,3}(?:[.\s]\d{3})*(?:[.,]\d{1,2})?|\d+(?:[.,]\d{1,2})?)', s)
    if not m:
        return None

    num = m.group(1).replace(" ", "")

    # normalize thousand/decimal separators
    if "," in num and "." in num:
        if num.rfind(",") > num.rfind("."):
            num = num.replace(".", "").replace(",", ".")
        else:
            num = num.replace(",", "")
    else:
        if "," in num:
            num = num.replace(".", "").replace(",", ".")

    # if user wrote "5" (no decimals), treat as "5.00"
    if re.fullmatch(r"\d+", num):
        num = num + ".00"
    elif re.fullmatch(r"\d+\.\d", num):
        num = num + "0"

    try:
        return Decimal(num).quantize(Decimal("0.01"))
    except InvalidOperation:
        return None


def _format_eur(d: Decimal) -> str:
    return f"{d:.2f}".replace(".", ",") + " EUR"

def _extract_amount_after_keyword(text: str, keywords: list[str]) -> Decimal | None:
    """
    Finds an amount near keywords, e.g. "Trinkgeld 5 €", "Tip: 4,00", "inkl. Trinkgeld 5 Euro".
    Returns Decimal or None.
    """
    if not text:
        return None
    t = text.upper()

    # patterns like: TRINKGELD 5 €, TIP: 4,00, INKL TRINKGELD 5 EURO
    # We capture a number with optional decimals and optional ",-" style.
    for kw in keywords:
        # allow up to ~25 chars between keyword and number (handles "inkl. ...")
        pat = rf"{kw}[^0-9]{{0,25}}(\d{{1,3}}(?:[.\s]\d{{3}})*(?:[.,]\d{{1,2}})?|\d+(?:[.,]\d{{1,2}})?)(?:\s*[-–—]?\s*)?(?:EUR|EURO|€)?"
        m = re.search(pat, t)
        if m:
            raw = m.group(1)
            # handle "5" -> "5.00" (we treat missing decimals as .00)
            amt = _parse_eur_amount(raw)
            if amt is not None:
                # if user wrote "5" without decimals, _parse_eur_amount returns 5.00 already
                return amt
    return None

def apply_tip_logic(bew_data: dict, ocr_text: str, email_text: str | None = None) -> dict:

    """
    Priority:
    1) Email: if it contains explicit total OR explicit tip -> override.
    2) OCR: if it contains explicit total (Gesamtbetrag, Zu zahlen) -> use it.
    3) OCR tip + base amount -> add.
    4) fallback to LLM amount.
    """
    out = bew_data.copy()

    base_amount = _parse_eur_amount(out.get("betrag", ""))

    # -------------------------
    # 1) EMAIL OVERRIDE
    # -------------------------
    em = (email_text or "")

    # If email provides an explicit total amount (e.g. "Gesamt 36,80", "insgesamt 36,80")
    email_total = _extract_amount_after_keyword(em, keywords=[
        "GESAMT", "INSGESAMT", "TOTAL", "ZU ZAHLEN", "SUMME"
    ])

    # If email provides explicit tip (e.g. "Trinkgeld 5 €")
    email_tip = _extract_amount_after_keyword(em, keywords=[
        "TRINKGELD", "TIP"
    ])

    if email_total is not None:
        out["betrag"] = _format_eur(email_total)
        out["betrag_quelle"] = "email_total"
        return out

    if email_tip is not None and base_amount is not None:
        out["betrag_rechnung"] = _format_eur(base_amount)
        out["trinkgeld"] = _format_eur(email_tip)
        out["betrag"] = _format_eur((base_amount + email_tip).quantize(Decimal("0.01")))
        out["betrag_quelle"] = "email_tip_plus_base"
        return out

    # -------------------------
    # 2) OCR TOTAL
    # -------------------------
    txt = (ocr_text or "").upper()

    total_patterns = [
        r"GESAMTBETRAG\s*[:\-]\s*([0-9\., ]+)",
        r"ZU\s*ZAHLEN\s*[:\-]\s*([0-9\., ]+)",
        r"TOTAL\s*[:\-]\s*([0-9\., ]+)",
        r"SUMME\s*[:\-]\s*([0-9\., ]+)",
        r"AMOUNT\s*DUE\s*[:\-]\s*([0-9\., ]+)",
    ]

    for pat in total_patterns:
        m = re.search(pat, txt)
        if m:
            total = _parse_eur_amount(m.group(1))
            if total is not None:
                out["betrag"] = _format_eur(total)
                out["betrag_quelle"] = "ocr_total"
                return out

    # -------------------------
    # 3) OCR TIP + BASE
    # -------------------------
    m_tip = re.search(r"(TIP\s*/\s*EXTRA|TRINKGELD|EXTRA)\s*[:\-]\s*([0-9\., ]+)", txt)
    if m_tip:
        tip = _parse_eur_amount(m_tip.group(2))
        if tip is not None and base_amount is not None:
            out["betrag_rechnung"] = _format_eur(base_amount)
            out["trinkgeld"] = _format_eur(tip)
            out["betrag"] = _format_eur((base_amount + tip).quantize(Decimal("0.01")))
            out["betrag_quelle"] = "ocr_tip_plus_base"
            return out

    out["betrag_quelle"] = "llm_amount"
    return out

# --------------------------------------------------
# PDFs mergen: Bon + Formular
# --------------------------------------------------

def merge_pdfs(receipt_path: str) -> None:
    """
    Merged den Bon (receipt_path) und das ausgefüllte Formular (OUTPUT_PDF_FORM)
    zu einer finalen PDF (OUTPUT_FINAL_PDF).
    """

    writer = PdfWriter()

    # 1) Bon-Seiten
    receipt_pdf = PdfReader(receipt_path)
    for page in receipt_pdf.pages:
        writer.add_page(page)

    # 2) Formular-Seiten
    form_pdf = PdfReader(OUTPUT_PDF_FORM)
    for page in form_pdf.pages:
        writer.add_page(page)

    with open(OUTPUT_FINAL_PDF, "wb") as f:
        writer.write(f)


# --------------------------------------------------
# Endpoint: /build-bewirtungsbeleg
# (Nimmt fertige Daten + Bon, baut PDF)
# --------------------------------------------------

@app.post("/build-bewirtungsbeleg")
async def build_bewirtungsbeleg(
    data: str = Form(...),
    receipt: UploadFile = File(...),
):
    """
    Erwartet:
    - data: JSON-String mit Bewirtungsdaten
    - receipt: Bon (PDF/JPG/PNG)

    Füllt das Template mit den Daten, konvertiert zu PDF
    und merged Bon + Formular zu einer finalen PDF.
    """

    bew_data = json.loads(data)

    # Bon speichern
    os.makedirs("input", exist_ok=True)
    receipt_path = os.path.join("input", receipt.filename)
    with open(receipt_path, "wb") as f:
        f.write(await receipt.read())

    # Formular erzeugen (DOCX -> PDF)
    generate_form_pdf(bew_data)

    # PDFs mergen
    merge_pdfs(receipt_path)

    return FileResponse(
        OUTPUT_FINAL_PDF,
        filename="bewirtungsbeleg_final.pdf",
        media_type="application/pdf",
    )


# --------------------------------------------------
# Endpoint: /full-agent
# (OCR + LLM + PDF, End-to-End)
# --------------------------------------------------

@app.post("/full-agent")
async def full_agent(
    email_text: str = Form(...),
    receipt: UploadFile = File(...),
):
    """
    Nimmt:
    - email_text: Text aus der E-Mail (Kontext, Anlass, Personen, etc.)
    - receipt: Bon (PDF/JPG/PNG)

    Pipeline:
    1. Bon speichern
    2. OCR mit ocr_bon
    3. Extraktion strukturierter Daten mit extract_bewirtungsbeleg_gemini
    4. Template füllen & Formular-PDF erzeugen
    5. Bon + Formular mergen
    6. Finale PDF zurückgeben
    """

    # 1) Bon speichern
    os.makedirs("input", exist_ok=True)
    receipt_path = os.path.join("input", receipt.filename)
    with open(receipt_path, "wb") as f:
        f.write(await receipt.read())

    # 2) OCR
    ocr_text = ocr_bon(receipt_path)
    print("----- OCR TEXT -----")
    print(ocr_text)
    print("----- END OCR -----")

    # 3) Strukturierte Daten (LLM)
    bew_data = extract_bewirtungsdaten_gemini(ocr_text, email_text)
    print("----- EMAIL TEXT START -----")
    print(repr(email_text[:500] if email_text else ""))
    print("----- EMAIL TEXT END -----")

    bew_data_before = bew_data.copy()

    bew_data = apply_tip_logic(bew_data, ocr_text=ocr_text, email_text=email_text)

    print("----- TIP LOGIC RESULT -----")
    print("betrag vorher:", bew_data_before.get("betrag"))
    print("betrag nachher:", bew_data.get("betrag"), "quelle:", bew_data.get("betrag_quelle"))
    print("trinkgeld:", bew_data.get("trinkgeld"), "betrag_rechnung:", bew_data.get("betrag_rechnung"))
    print("----- END TIP LOGIC RESULT -----")

    print("----- EXTRACTED DATA -----")
    print(json.dumps(bew_data, indent=2, ensure_ascii=False))
    print("----- END DATA -----")


    # 4) Formular erzeugen
    generate_form_pdf(bew_data)

    # 5) Bon + Formular mergen
    merge_pdfs(receipt_path)


    # 6) PDF mit sauberem Dateinamen zurückgeben
    def safe(s: str) -> str:
        return (
            s.replace(" ", "-")
            .replace("/", "-")
            .replace("\\", "-")
            .replace("€", "EUR")
        )

    restaurant = safe(bew_data.get("restaurant", "Restaurant"))
    betrag = safe(bew_data.get("betrag", ""))
    datum = date.today().strftime("%Y-%m-%d")

    filename = f"Bewirtungsbeleg_{datum}_{restaurant}_{betrag}.pdf"

    return FileResponse(
        OUTPUT_FINAL_PDF,
        filename=filename,
        media_type="application/pdf",
    )
